"""Build the one-page lecture abstract (Word document) for the professors.

Usage:  python build_abstract.py
Output: ../nus_lecture_abstract.docx  (repo root, next to the slide deck)
"""

from pathlib import Path

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Pt, RGBColor

NAVY = RGBColor(0x1F, 0x38, 0x64)
GREY = RGBColor(0x5A, 0x5A, 0x5A)

SPEAKER = "Wanghuan Li"
AFFILIATION = ("Software Engineer @ TikTok (AI) · "
               "Part-time PhD Student, NTU Singapore")
CONTACT = "liwanghuan43@gmail.com"

TITLE = ("AI Agents for Retail Analytics: RFM Customer Segmentation "
         "& Automated Reporting")
COURSE = ("Guest Lecture · Data Analytics in Business · "
          "National University of Singapore · 2 hours")

BIO = (
    f"{SPEAKER} is a software engineer at TikTok focused on AI development, "
    "with more than seven years of industry experience. At TikTok, this work "
    "centres on building AI agents that autonomously clean data, run "
    "analyses, and write reports for business stakeholders. Alongside "
    f"industry work, {SPEAKER} is a part-time PhD student at Nanyang "
    "Technological University (NTU), Singapore, with research interests in "
    "LLM-based agents and applied data analytics."
)

ABSTRACT = [
    "This two-hour guest lecture gives students a faithful, end-to-end "
    "experience of a real industry scenario: turning raw, messy retail "
    "transaction data into an actionable marketing strategy, fully "
    "automated by an AI agent. Rather than a textbook exercise, every "
    "element of the session — the dirty dataset, the repeated monthly "
    "analysis, the plain-language report for stakeholders — mirrors what "
    "data analytics teams at technology and retail companies actually do.",

    "The lecture first develops the business case for customer "
    "segmentation and the fundamentals of RFM (Recency, Frequency, "
    "Monetary) analysis, the most widely used behavioural segmentation "
    "method in retail CRM. It then shows how a manual analyst workflow "
    "becomes an AI agent: a five-step pipeline (ingest → clean → score → "
    "segment → report) that is deterministic where correctness matters "
    "and uses a large language model where judgement and language matter. "
    "A live demonstration runs the agent on a deliberately dirty dataset "
    "of 5,500 transactions, producing interactive dashboards and a "
    "downloadable strategy report. The session closes with per-segment "
    "marketing strategies and KPIs, A/B-test-based validation, and a "
    "discussion of the limitations and ethics of both RFM and "
    "LLM-in-the-loop analytics. All code, slides, and data are shared "
    "with students, who can rebuild the entire agent themselves.",
]


def add_heading(doc, text):
    p = doc.add_paragraph()
    run = p.add_run(text)
    run.font.size, run.font.bold, run.font.color.rgb = Pt(14), True, NAVY
    p.space_before, p.space_after = Pt(14), Pt(4)
    return p


def add_body(doc, text):
    p = doc.add_paragraph()
    run = p.add_run(text)
    run.font.size = Pt(11)
    p.paragraph_format.space_after = Pt(8)
    p.paragraph_format.line_spacing = 1.25
    return p


doc = Document()

# Title block
p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = p.add_run(TITLE)
run.font.size, run.font.bold, run.font.color.rgb = Pt(18), True, NAVY

p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = p.add_run(COURSE)
run.font.size, run.font.color.rgb = Pt(11), GREY

p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = p.add_run(f"{SPEAKER}  ·  {AFFILIATION}  ·  {CONTACT}")
run.font.size, run.font.italic, run.font.color.rgb = Pt(10), True, GREY
p.paragraph_format.space_after = Pt(12)

add_heading(doc, "Speaker Bio")
add_body(doc, BIO)

add_heading(doc, "Abstract")
for para in ABSTRACT:
    add_body(doc, para)

out = Path(__file__).parent.parent / "nus_lecture_abstract.docx"
doc.save(out)
print(f"Saved -> {out}")
